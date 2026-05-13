from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


ROOT = Path(__file__).resolve().parent
REFERENCE = ROOT / "aw_textbook_clean_reference.docx"
SPECIMEN_MD = ROOT / "aw_textbook_style_specimen.md"
SPECIMEN_DOCX = ROOT / "aw_textbook_style_specimen.docx"
LUA_FILTER = ROOT / "aw_textbook_div_styles.lua"

FONTS = {
    "body": "Congenial",
    "heading": "Roboto Condensed",
    "heading_medium": "Roboto Condensed Medium",
    "mono": "Consolas",
}

COLORS = {
    "charcoal": "2D2D2D",
    "near_black": "232323",
    "grey": "6E6E6E",
    "muted_blue": "3F6EB4",
    "blue_grey": "56768C",
    "muted_green": "528F4A",
    "muted_red": "B03E3E",
    "muted_orange": "C47D2A",
    "soft_grey": "969696",
    "dark_navy_grey": "2D4155",
    "medium_blue_grey": "46647D",
    "very_light_blue": "F2F7FC",
    "very_light_grey": "F7F7F7",
    "very_light_green": "F4FAF2",
    "very_light_red": "FCF4F4",
    "very_light_orange": "FDF8F0",
    "border_grey": "CDD2D6",
}


def twips_pt(value: float) -> str:
    return str(round(value * 20))


def twips_cm(value: float) -> str:
    return str(round(value * 567))


def half_points(value: float) -> str:
    return str(round(value * 2))


def ensure_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_children(parent, tags):
    for tag in tags:
        for child in list(parent.findall(qn(tag))):
            parent.remove(child)


def set_fonts(rpr, font_name):
    remove_children(rpr, ["w:rFonts"])
    rfonts = OxmlElement("w:rFonts")
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)
    rpr.append(rfonts)


def set_text_props(style, font_name, size_pt, color, italic=False, bold=False):
    rpr = ensure_child(style.element, "w:rPr")
    set_fonts(rpr, font_name)
    remove_children(rpr, ["w:sz", "w:szCs", "w:color", "w:i", "w:b"])
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), half_points(size_pt))
        rpr.append(el)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rpr.append(color_el)
    if italic:
        rpr.append(OxmlElement("w:i"))
    b = OxmlElement("w:b")
    b.set(qn("w:val"), "1" if bold else "0")
    rpr.append(b)


def set_paragraph_props(
    style,
    *,
    before=0,
    after=0,
    line=1.0,
    left_cm=0,
    right_cm=0,
    hanging_cm=0,
    align="left",
    keep_next=False,
    keep_lines=False,
    page_break_before=False,
    widow=True,
    shading=None,
    borders=None,
):
    ppr = ensure_child(style.element, "w:pPr")
    remove_children(
        ppr,
        [
            "w:spacing",
            "w:ind",
            "w:jc",
            "w:keepNext",
            "w:keepLines",
            "w:pageBreakBefore",
            "w:widowControl",
            "w:shd",
            "w:pBdr",
        ],
    )
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), twips_pt(before))
    spacing.set(qn("w:after"), twips_pt(after))
    spacing.set(qn("w:line"), str(round(line * 240)))
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), twips_cm(left_cm))
    ind.set(qn("w:right"), twips_cm(right_cm))
    if hanging_cm:
        ind.set(qn("w:hanging"), twips_cm(hanging_cm))
    ppr.append(ind)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), align)
    ppr.append(jc)
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))
    if page_break_before:
        ppr.append(OxmlElement("w:pageBreakBefore"))
    if widow:
        ppr.append(OxmlElement("w:widowControl"))
    if shading:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), shading)
        ppr.append(shd)
    if borders:
        pbdr = OxmlElement("w:pBdr")
        for side, spec in borders.items():
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), spec.get("val", "single"))
            border.set(qn("w:sz"), str(round(spec["pt"] * 8)))
            border.set(qn("w:space"), str(spec.get("space", 1)))
            border.set(qn("w:color"), spec["color"])
            pbdr.append(border)
        ppr.append(pbdr)


def get_or_add_style(styles, name, style_type, base=None, next_style=None):
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, style_type)
    if base:
        style.base_style = styles[base]
    if next_style and style_type == WD_STYLE_TYPE.PARAGRAPH:
        style.next_paragraph_style = styles[next_style]
    return style


def create_linked_character_style(styles, para_style_name, font_name, size, color, italic=False, bold=False):
    char_name = f"{para_style_name} Char"
    char_style = get_or_add_style(styles, char_name, WD_STYLE_TYPE.CHARACTER)
    set_text_props(char_style, font_name, size, color, italic=italic, bold=bold)
    para_style = styles[para_style_name]
    para_el = para_style.element
    char_el = char_style.element
    remove_children(para_el, ["w:link"])
    remove_children(char_el, ["w:link"])
    p_link = OxmlElement("w:link")
    p_link.set(qn("w:val"), char_style.style_id)
    para_el.append(p_link)
    c_link = OxmlElement("w:link")
    c_link.set(qn("w:val"), para_style.style_id)
    char_el.append(c_link)


def define_paragraph_styles(doc: Document):
    styles = doc.styles
    specs = [
        ("Normal", None, "Body Text", FONTS["body"], 10.5, COLORS["charcoal"], 0, 6, 1.30, 0, 0, False, False, False, None, None, False, False),
        ("Body Text", "Normal", "Body Text", FONTS["body"], 10.5, COLORS["charcoal"], 0, 7, 1.30, 0, 0, False, False, False, None, None, True, False),
        ("Body Text Compact", "Body Text", "Body Text", FONTS["body"], 10, COLORS["charcoal"], 0, 4, 1.15, 0, 0, False, False, False, None, None, True, False),
        ("Heading 1", "Normal", "Body Text", FONTS["heading_medium"], 24, "FFFFFF", 0, 16, 1.0, 0, 0, True, True, True, COLORS["dark_navy_grey"], None, True, False),
        ("Heading 2", "Normal", "Body Text", FONTS["heading_medium"], 17, "FFFFFF", 20, 10, 1.0, 0, 0, True, True, False, COLORS["medium_blue_grey"], None, True, False),
        ("Heading 3", "Normal", "Body Text", FONTS["heading"], 14, COLORS["near_black"], 18, 6, 1.0, 0.3, 0, True, True, False, "EDF4FA", {"left": {"pt": 4.5, "color": COLORS["blue_grey"], "space": 3}}, True, False),
        ("Heading 4", "Normal", "Body Text", FONTS["heading"], 12, COLORS["charcoal"], 14, 5, 1.0, 0, 0, True, True, False, None, {"bottom": {"pt": 0.75, "color": COLORS["border_grey"], "space": 1}}, True, False),
        ("Heading 5", "Normal", "Body Text", FONTS["heading"], 11, COLORS["charcoal"], 8, 4, 1.0, 0, 0, True, True, False, None, {"bottom": {"pt": 0.5, "color": COLORS["border_grey"], "space": 1}}, True, False),
        ("Homework Target", "Body Text", "Body Text", FONTS["body"], 10, COLORS["grey"], 0, 12, 1.15, 0, 0, True, False, False, None, None, True, True),
        ("Reflection Text", "Body Text", "Body Text", FONTS["body"], 10, COLORS["grey"], 6, 6, 1.20, 0.5, 0, False, False, False, None, None, True, True),
        ("After List", "Body Text", "Body Text", FONTS["body"], 10.5, COLORS["charcoal"], 8, 6, 1.30, 0, 0, False, False, False, None, None, True, False),
        ("Caption", "Body Text", "Body Text", FONTS["body"], 9, COLORS["grey"], 4, 8, 1.10, 0, 0, False, False, False, None, None, True, True),
        ("Page Header", "Normal", "Page Header", FONTS["body"], 8, COLORS["grey"], 0, 0, 1.0, 0, 0, False, False, False, None, None, True, False),
        ("Page Footer", "Normal", "Page Footer", FONTS["body"], 8, COLORS["grey"], 0, 0, 1.0, 0, 0, False, False, False, None, None, True, False),
        ("List Bullet", "Body Text", "Body Text", FONTS["body"], 10.5, COLORS["charcoal"], 0, 3, 1.15, 0.8, 0, False, False, False, None, None, True, False),
        ("List Bullet 2", "Body Text", "Body Text", FONTS["body"], 10, COLORS["charcoal"], 0, 2, 1.10, 1.3, 0, False, False, False, None, None, True, False),
        ("List Number", "Body Text", "Body Text", FONTS["body"], 10.5, COLORS["charcoal"], 0, 3, 1.15, 0.9, 0, False, False, False, None, None, True, False),
        ("List Number 2", "Body Text", "Body Text", FONTS["body"], 10, COLORS["charcoal"], 0, 2, 1.10, 1.4, 0, False, False, False, None, None, True, False),
        ("Checklist", "Body Text", "Body Text", FONTS["body"], 10.5, COLORS["charcoal"], 0, 4, 1.15, 0.9, 0, False, False, False, None, None, True, False),
    ]
    for name, base, next_style, font, size, color, before, after, line, left, right, keep_next, keep_lines, page_break, shading, borders, linked, italic in specs:
        style = get_or_add_style(styles, name, WD_STYLE_TYPE.PARAGRAPH, base, next_style)
        set_text_props(style, font, size, color, italic=italic, bold=False)
        set_paragraph_props(
            style,
            before=before,
            after=after,
            line=line,
            left_cm=left,
            right_cm=right,
            keep_next=keep_next,
            keep_lines=keep_lines,
            page_break_before=page_break,
            shading=shading,
            borders=borders,
        )
        if linked and name != "Normal":
            create_linked_character_style(styles, name, font, size, color, italic=italic, bold=False)


def define_semantic_styles(doc: Document):
    specs = [
        ("Learn Process", COLORS["muted_blue"], COLORS["very_light_blue"], 4.5, 10, 10, 0.35, 0, 1.20, False),
        ("Learn Language", COLORS["blue_grey"], "F0F5F8", 4.5, 10, 10, 0.35, 0, 1.20, False),
        ("Learn Principle", COLORS["muted_blue"], "EAF3FC", 6, 10, 10, 0.35, 0, 1.20, False),
        ("Learn Transfer", COLORS["muted_orange"], COLORS["very_light_orange"], 6, 10, 10, 0.35, 0, 1.20, False),
        ("Learn Teaching", COLORS["muted_blue"], COLORS["very_light_blue"], 4.5, 10, 10, 0.35, 0, 1.20, False),
        ("Learn Note", COLORS["soft_grey"], COLORS["very_light_grey"], 3, 8, 8, 0.35, 0, 1.20, False),
        ("Model", COLORS["blue_grey"], "F7FAFC", 4.5, 10, 10, 0.35, 0.35, 1.20, False),
        ("Model Bad", COLORS["muted_red"], COLORS["very_light_red"], 6, 10, 10, 0.5, 0.5, 1.20, False),
        ("Model Good", COLORS["muted_green"], COLORS["very_light_green"], 6, 10, 10, 0.5, 0.5, 1.20, False),
        ("Worked Example", COLORS["blue_grey"], "F0F5F8", 4.5, 10, 10, 0.35, 0, 1.20, False),
        ("Example", COLORS["soft_grey"], COLORS["very_light_grey"], 3, 8, 8, 0.35, 0, 1.20, False),
        ("Placeholder", COLORS["soft_grey"], None, 0.5, 8, 8, 0, 0, 1.15, True),
        ("Self Study", COLORS["soft_grey"], COLORS["very_light_grey"], 3, 8, 8, 0.35, 0, 1.20, False),
        ("Annotation", COLORS["soft_grey"], None, 3, 4, 6, 0.35, 0, 1.15, False),
    ]
    for name, border_color, fill, border_pt, before, after, left, right, line, italic in specs:
        style = get_or_add_style(doc.styles, name, WD_STYLE_TYPE.PARAGRAPH, "Body Text", "Body Text")
        size = 9.5 if name == "Annotation" else 10
        color = COLORS["grey"] if name in {"Placeholder", "Annotation"} else COLORS["charcoal"]
        set_text_props(style, FONTS["body"], size, color, italic=italic, bold=False)
        if name in {"Model Bad", "Model Good"}:
            borders = {
                "top": {"pt": 0.75, "color": border_color, "space": 1},
                "left": {"pt": border_pt, "color": border_color, "space": 2},
                "bottom": {"pt": 0.75, "color": border_color, "space": 1},
                "right": {"pt": 0.75, "color": border_color, "space": 1},
            }
        elif name == "Placeholder":
            borders = {
                side: {"pt": border_pt, "color": border_color, "space": 2, "val": "dotted"}
                for side in ("top", "left", "bottom", "right")
            }
        else:
            borders = {"left": {"pt": border_pt, "color": border_color, "space": 3}}
        set_paragraph_props(
            style,
            before=before,
            after=after,
            line=line,
            left_cm=left,
            right_cm=right,
            shading=fill,
            borders=borders,
        )
        create_linked_character_style(doc.styles, name, FONTS["body"], size, color, italic=italic, bold=False)


def define_character_styles(doc: Document):
    specs = [
        ("Strong", FONTS["body"], 10.5, COLORS["charcoal"], False, False),
        ("Emphasis", FONTS["body"], 10.5, COLORS["charcoal"], True, False),
        ("Metadata Code", FONTS["heading"], 8, COLORS["grey"], False, False),
        ("Activity Star", FONTS["heading"], 8, COLORS["muted_orange"], False, False),
        ("Learn Label", FONTS["heading"], 10, COLORS["muted_blue"], False, False),
        ("Model Label Bad", FONTS["heading_medium"], 10, COLORS["muted_red"], False, False),
        ("Model Label Good", FONTS["heading_medium"], 10, COLORS["muted_green"], False, False),
        ("Inline Code", FONTS["mono"], 9, COLORS["charcoal"], False, False),
        ("Placeholder Token", FONTS["mono"], 8, COLORS["grey"], False, False),
        ("Small Note", FONTS["body"], 9, COLORS["grey"], True, False),
    ]
    for name, font, size, color, italic, bold in specs:
        style = get_or_add_style(doc.styles, name, WD_STYLE_TYPE.CHARACTER)
        set_text_props(style, font, size, color, italic=italic, bold=bold)


def add_field(run, instr):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    run._r.append(instr_el)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.gutter = Cm(0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    section.different_first_page_header_footer = True

    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))

    for footer, even in ((section.footer, False), (section.even_page_footer, True)):
        p = footer.paragraphs[0]
        p.style = doc.styles["Page Footer"]
        p.alignment = 0
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15))
        if even:
            add_field(p.add_run(), " PAGE ")
            p.add_run("\tAdministrative Writing, Advanced")
        else:
            p.add_run("Short unit title\t")
            add_field(p.add_run(), " PAGE ")
    section.first_page_footer.paragraphs[0].style = doc.styles["Page Footer"]


def write_lua_filter():
    mapping = {
        "learn-process": "Learn Process",
        "learn-language": "Learn Language",
        "learn-principle": "Learn Principle",
        "learn-transfer": "Learn Transfer",
        "learn-teaching": "Learn Teaching",
        "learn-note": "Learn Note",
        "model": "Model",
        "model-bad": "Model Bad",
        "model-good": "Model Good",
        "worked-example": "Worked Example",
        "example": "Example",
        "placeholder": "Placeholder",
        "self-study": "Self Study",
        "annotation": "Annotation",
    }
    lines = [
        "-- Auto-generated by build_aw_textbook_reference.py",
        "-- Maps Step 1 semantic Div classes to Word custom styles for Pandoc DOCX output.",
        "local style_for_class = {",
    ]
    for k, v in mapping.items():
        lines.append(f'  ["{k}"] = "{v}",')
    lines += [
        "}",
        "",
        "-- The book source uses standalone --- as combined-draft separators.",
        "-- They must not become Word page breaks or visible horizontal rules.",
        "function HorizontalRule(el)",
        "  return {}",
        "end",
        "",
        "-- Avoid Pandoc's direct bold/italic output overriding the lighter Word style system.",
        "function Strong(el)",
        "  return pandoc.Span(el.content, pandoc.Attr('', {}, {['custom-style'] = 'Strong'}))",
        "end",
        "",
        "function Emph(el)",
        "  return pandoc.Span(el.content, pandoc.Attr('', {}, {['custom-style'] = 'Emphasis'}))",
        "end",
        "",
        "local function preserve_div_line_breaks(el)",
        "  return pandoc.walk_block(el, {",
        "    SoftBreak = function(_)",
        "      return pandoc.LineBreak()",
        "    end",
        "  })",
        "end",
        "",
        "function Div(el)",
        "  for _, class in ipairs(el.classes) do",
        "    local style = style_for_class[class]",
        "    if style then",
        "      el.attributes['custom-style'] = style",
        "      return preserve_div_line_breaks(el)",
        "    end",
        "  end",
        "  return el",
        "end",
        "",
    ]
    LUA_FILTER.write_text("\n".join(lines), encoding="utf-8")


def patch_styles_xml(docx_path: Path):
    tmp = Path(tempfile.mkdtemp())
    try:
        with ZipFile(docx_path) as zin:
            zin.extractall(tmp)
        styles_path = tmp / "word" / "styles.xml"
        numbering_path = tmp / "word" / "numbering.xml"
        tree = etree.parse(str(styles_path))
        root = tree.getroot()
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # Remove bold from all reduced-weight heading/table styles explicitly.
        for style_id in ("Heading1", "Heading2", "TableHeader", "LightShading-Accent1"):
            for style in root.xpath(f'.//w:style[@w:styleId="{style_id}"]', namespaces=ns):
                rpr = style.find("w:rPr", ns)
                if rpr is None:
                    rpr = etree.SubElement(style, qn("w:rPr"))
                for b in rpr.findall("w:b", ns):
                    rpr.remove(b)
                off = etree.SubElement(rpr, qn("w:b"))
                off.set(qn("w:val"), "0")

        # Add table styles not fully representable through python-docx.
        for name, style_id, fill, text_color in [
            ("AW Standard Table", "AWStandardTable", COLORS["dark_navy_grey"], "FFFFFF"),
            ("AW Rubric Table", "AWRubricTable", COLORS["medium_blue_grey"], "FFFFFF"),
            ("AW Phrase Bank Table", "AWPhraseBankTable", COLORS["very_light_blue"], COLORS["near_black"]),
            ("AW Comparison Table", "AWComparisonTable", COLORS["very_light_grey"], COLORS["near_black"]),
        ]:
            old = root.xpath(f'.//w:style[@w:styleId="{style_id}"]', namespaces=ns)
            for node in old:
                root.remove(node)
            style = etree.SubElement(root, qn("w:style"))
            style.set(qn("w:type"), "table")
            style.set(qn("w:customStyle"), "1")
            style.set(qn("w:styleId"), style_id)
            etree.SubElement(style, qn("w:name")).set(qn("w:val"), name)
            tbl_pr = etree.SubElement(style, qn("w:tblPr"))
            jc = etree.SubElement(tbl_pr, qn("w:jc"))
            jc.set(qn("w:val"), "center")
            layout = etree.SubElement(tbl_pr, qn("w:tblLayout"))
            layout.set(qn("w:type"), "fixed")
            borders = etree.SubElement(tbl_pr, qn("w:tblBorders"))
            for side in ("top", "left", "bottom", "right"):
                b = etree.SubElement(borders, qn(f"w:{side}"))
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "6")
                b.set(qn("w:color"), COLORS["soft_grey"])
            for side in ("insideH", "insideV"):
                b = etree.SubElement(borders, qn(f"w:{side}"))
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), COLORS["border_grey"])
            cell_mar = etree.SubElement(tbl_pr, qn("w:tblCellMar"))
            for side, pt in {"top": 5, "bottom": 5, "left": 6, "right": 6}.items():
                mar = etree.SubElement(cell_mar, qn(f"w:{side}"))
                mar.set(qn("w:w"), twips_pt(pt))
                mar.set(qn("w:type"), "dxa")
            whole = etree.SubElement(style, qn("w:tblStylePr"))
            whole.set(qn("w:type"), "wholeTable")
            ppr = etree.SubElement(whole, qn("w:pPr"))
            spacing = etree.SubElement(ppr, qn("w:spacing"))
            spacing.set(qn("w:after"), twips_pt(2))
            spacing.set(qn("w:line"), str(round(1.10 * 240)))
            spacing.set(qn("w:lineRule"), "auto")
            rpr = etree.SubElement(whole, qn("w:rPr"))
            etree.SubElement(rpr, qn("w:b")).set(qn("w:val"), "0")
            etree.SubElement(rpr, qn("w:sz")).set(qn("w:val"), half_points(9.5))
            etree.SubElement(rpr, qn("w:color")).set(qn("w:val"), COLORS["charcoal"])
            first = etree.SubElement(style, qn("w:tblStylePr"))
            first.set(qn("w:type"), "firstRow")
            first_rpr = etree.SubElement(first, qn("w:rPr"))
            etree.SubElement(first_rpr, qn("w:b")).set(qn("w:val"), "0")
            etree.SubElement(first_rpr, qn("w:color")).set(qn("w:val"), text_color)
            first_tcpr = etree.SubElement(first, qn("w:tcPr"))
            shd = etree.SubElement(first_tcpr, qn("w:shd"))
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), fill)

        tree.write(str(styles_path), encoding="UTF-8", xml_declaration=True, standalone=True)

        if numbering_path.exists():
            n_tree = etree.parse(str(numbering_path))
            n_root = n_tree.getroot()
            # Keep numbering templates stable and explicit. Pandoc may add numIds, but styles are linked here.
            n_tree.write(str(numbering_path), encoding="UTF-8", xml_declaration=True, standalone=True)

        rebuilt = docx_path.with_suffix(".rebuilt.docx")
        with ZipFile(rebuilt, "w", ZIP_DEFLATED) as zout:
            for path in tmp.rglob("*"):
                if path.is_file():
                    zout.write(path, path.relative_to(tmp).as_posix())
        rebuilt.replace(docx_path)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def build_reference():
    doc = Document()
    define_paragraph_styles(doc)
    define_semantic_styles(doc)
    define_character_styles(doc)
    configure_document(doc)
    doc.core_properties.title = "Administrative Writing Textbook Clean Reference"
    doc.core_properties.subject = "Pandoc DOCX reference styles"
    doc.save(REFERENCE)
    patch_styles_xml(REFERENCE)
    write_lua_filter()


def build_specimen():
    md = """# Administrative Writing, Advanced

## Unit Title Sample

Body text sample paragraph for reading rhythm and after-spacing.

### Section Sample

---

#### Activity Sample — Subtitle (B1) ★

##### Micro-Theory Heading

- Bullet item one
- Bullet item two
  - Nested bullet item

1. Number item one
1. Number item two
   1. Nested number item

::: learn-process
**Why This Works:** A short process explanation shows the semantic learning block spacing.
:::

::: learn-language
**Language Note:** A short language explanation shows the blue-grey learning block.
:::

::: learn-principle
**Key Principle:** A major concept uses stronger emphasis.
:::

::: learn-transfer
**Transfer Reminder:** A transfer reminder uses the orange accent.
:::

::: learn-teaching
**Teaching Point:** A teaching note uses the learning accent.
:::

::: learn-note
**Note:** A lower-intensity note uses a quiet grey accent.
:::

::: model-bad
**Original Text — Too Direct**

We need this today. Send it immediately.
:::

::: model-good
**Revised Text — Diplomatic**

Could you please send this by the end of today if possible?
:::

::: worked-example
**Worked Example:** First identify the reader, then choose the action, then soften the request.
:::

::: example
**Example:** This is a simple example block.
:::

::: self-study
**Self-Study:** Read the message from the recipient's point of view.
:::

::: annotation
**Annotation:** This phrase reduces pressure while keeping the request clear.
:::

::: placeholder
{{PH-SPECIMEN-001}}
:::

| Column A | Column B |
|---|---|
| Sample | Table body |
| Sample | Table body |
"""
    SPECIMEN_MD.write_text(md, encoding="utf-8")
    cmd = [
        "pandoc",
        str(SPECIMEN_MD),
        "--from",
        "markdown-yaml_metadata_block",
        "--to",
        "docx",
        "--reference-doc",
        str(REFERENCE),
        "--lua-filter",
        str(LUA_FILTER),
        "--output",
        str(SPECIMEN_DOCX),
    ]
    subprocess.run(cmd, check=True)


def main():
    build_reference()
    build_specimen()
    print(f"Wrote {REFERENCE}")
    print(f"Wrote {LUA_FILTER}")
    print(f"Wrote {SPECIMEN_MD}")
    print(f"Wrote {SPECIMEN_DOCX}")


if __name__ == "__main__":
    main()
